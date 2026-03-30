#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

from docx import Document
from sklearn.manifold import MDS

ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / 'data'
ASSETS_DIR = ROOT / 'assets'
DOCS_DIR = ROOT / 'docs'
REAL_SAS_DIR = ROOT / 'source_materials' / 'real_sas_docs'
CONFIG_DIR = ROOT / 'config'
OVERRIDES_PATH = CONFIG_DIR / 'profile_overrides.json'
VERSION_TEXT = 'v1.6.1-fix-multi-click-after-react'

STOPWORDS = {
    'the','and','for','with','from','that','this','into','through','over','across','their','them','will','then','than','using','used','use',
    'onto','while','where','which','what','when','able','such','like','also','very','more','less','same','different','provide','provides',
    'support','supports','help','helps','allow','allows','application','solution','platform','business','global','regional','region','teams','team',
    'project','section','document','architecture','system','data','user','users','process','processes','model','models','service','services'
}

SECTION_HEADINGS = [
    'Background',
    'Business Needs and Benefits',
    'Objectives and Scope',
    'Assumptions',
    'Out of scope',
    'Technology Solution Governance (TSG)',
    'Business Capability View',
    'User Accessibility',
    'Supported Languages',
    'Application Monitoring',
    'Business Volume Projections',
    'Performance',
    'Availability Hours',
    'Compliance',
    'System Failure Conditions',
    'Licences',
    'Website Style',
    'Logical Architecture',
    'System Topology/Context',
    'Interface / Integration Item',
    'Architecture Layers',
    'Solution Services',
    'Introduction',
    'Data Dependencies – inputs',
    'Data Model',
    'System-of-record',
    'Data Archival',
    'ML and Analytics',
    'ML and Analytics (LLM Models included)',
    'Cloud Infrastructure/Services',
    'Cloud Software',
    'Environments and Hosting Location',
    'Security Questionnaire',
]

FAMILY_COLORS = {
    'Fraud & Entity Intelligence': '#2f6df6',
    'Pricing & Portfolio Optimization': '#8a3ffc',
    'Assistants & Document AI': '#119c92',
    'Other AI Application': '#64748b',
}

TECH_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'\bAKS\b|Azure Kubernetes(?: Service)?|Azure Kubernetes\b', 'AKS'),
    (r'Azure Databricks|Databricks\b', 'Azure Databricks'),
    (r'\bADLS\b|Azure Data Lake Storage', 'ADLS'),
    (r'Elasticsearch|\bELK\b', 'Elasticsearch'),
    (r'PostgreSQL|Postgres(?: Flexi Server)?', 'PostgreSQL'),
    (r'Azure Active Directory|\bAAD\b|Microsoft Entra ID|Azure AD', 'Azure Active Directory'),
    (r'Azure Key Vault|\bAKV\b', 'Azure Key Vault'),
    (r'Azure Log Analytics|Log Analytics', 'Azure Log Analytics'),
    (r'\bADF\b|Azure Data Factory', 'ADF'),
    (r'Cosmos(?: DB)?', 'Cosmos DB'),
    (r'Azure AI Search|AI Search|vector db', 'Azure AI Search'),
    (r'Azure Open ?AI|Open AI GPT|OpenAI', 'Azure OpenAI'),
    (r'Azure Blob Storage|Blob Storage', 'Azure Blob Storage'),
    (r'\bAPIM\b|API Management', 'APIM'),
    (r'\bRedis\b', 'Redis'),
    (r'AngularJS|Angular\b', 'AngularJS'),
    (r'NodeJS|Node\.js|Node JS', 'NodeJS'),
    (r'TypeScript', 'TypeScript'),
    (r'\bPython(?:\s*\d+(?:\.\d+)?)?\b', 'Python'),
    (r'FastAPI', 'FastAPI'),
    (r'Power BI', 'Power BI'),
    (r'Confluent|Kafka', 'Confluent Kafka'),
    (r'Snowflake', 'Snowflake'),
    (r'Azure ML|Machine Learning', 'Azure ML'),
    (r'Quantexa', 'Quantexa'),
    (r'OAuth2(?:\.0)?|OAuth', 'OAuth'),
]

CAPABILITY_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'question answering|q\s*&\s*a|q&a functionality|fast, accurate information retrieval', 'document Q&A'),
    (r'semantic search', 'semantic search'),
    (r'vector (?:search|embedding)', 'vector search'),
    (r'document upload|document ingestion|ingest(?:ion)? of documents', 'document ingestion'),
    (r'metadata extract|extract the metadata', 'metadata extraction'),
    (r'document comparison', 'document comparison'),
    (r'entity resolution', 'entity resolution'),
    (r'network analytics|relationship analytics|network visualization', 'network analytics'),
    (r'fraud detection', 'fraud detection'),
    (r'relationship search', 'relationship search'),
    (r'investigation', 'investigation workflows'),
    (r'risk assessment', 'risk assessment'),
    (r'portfolio monitoring|real[- ]time monitoring', 'portfolio monitoring'),
    (r'price optimization|pricing strateg|pricing decisions|dynamic pricing', 'price optimization'),
    (r'simulation', 'simulation'),
    (r'scenario testing', 'scenario testing'),
    (r'anomal(?:y|ies) detection|outlier', 'anomaly detection'),
    (r'decision support|decision-making|actionable insights', 'decision support'),
    (r'dashboard', 'dashboarding'),
    (r'model monitoring|data drift', 'model monitoring'),
    (r'semantic retrieval|retrieval augmented|rag', 'retrieval / RAG'),
    (r'catalog', 'application catalog'),
    (r'summarization', 'summarization'),
]

FEATURE_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'document upload|upload documents', 'document upload'),
    (r'question answering|q\s*&\s*a', 'question answering'),
    (r'chat history|prompts/threads|prompts and threads', 'chat history'),
    (r'user preferences|user profile', 'user preferences'),
    (r'app store|catalog', 'app store catalog'),
    (r'document comparison', 'document comparison'),
    (r'semantic retrieval', 'semantic retrieval'),
    (r'office plugin|ms office|outlook/word/ppt', 'office plugin access'),
    (r'entity graph', 'entity graph'),
    (r'network visualization', 'network visualization'),
    (r'dynamic entity resolution', 'dynamic entity resolution'),
    (r'task creation|tasks', 'task creation'),
    (r'alert thresholds|alerts', 'alert thresholds'),
    (r'investigator ui|quantexa ui', 'investigator UI'),
    (r'sso access|single sign', 'SSO access'),
    (r'monitoring pipeline|monitoring process', 'monitoring pipeline'),
    (r'config(?:uration)? upload|upload configuration', 'configuration upload'),
    (r'simulation', 'simulation library'),
    (r'power bi|dashboard', 'embedded Power BI dashboards'),
    (r'model monitoring', 'model monitoring'),
    (r'glossary', 'glossary'),
    (r'report generation|summary reports', 'report generation'),
]

DATA_DOMAIN_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'underwriting guidelines', 'underwriting guidelines'),
    (r'risk engineering documents?', 'risk engineering documents'),
    (r'legal docs?|legal documents?', 'legal documents'),
    (r'metadata', 'metadata'),
    (r'prompts?/threads?|threads?', 'prompts and threads'),
    (r'user profile', 'user profile'),
    (r'\bindividual\b', 'individual'),
    (r'\bbusiness\b', 'business'),
    (r'address', 'address'),
    (r'telephone|phone', 'telephone'),
    (r'email', 'email'),
    (r'documents?', 'documents'),
    (r'alerts?', 'alerts'),
    (r'tasks?', 'tasks'),
    (r'\bpolicy\b', 'policy'),
    (r'\bclaims?\b', 'claims'),
    (r'reference data', 'reference data'),
    (r'pricing', 'pricing'),
    (r'submissions?', 'submissions'),
    (r'technical premium', 'technical premium'),
    (r'loss cost', 'loss cost'),
    (r'portfolio metadata|portfolio data', 'portfolio metadata'),
]

MODEL_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'GPT-4o Mini|GPT 4o Mini|4o mini', 'GPT-4o Mini'),
    (r'Azure Document Intelligence|Doc Parser|Document Intelligence', 'Azure Document Intelligence'),
    (r'text-embedding-3-small|embedding model', 'text-embedding-3-small'),
    (r'entity resolution', 'entity resolution'),
    (r'network fraud model|fraud model', 'network fraud model'),
    (r'knowledge graph', 'knowledge graph'),
    (r'conversion model|L0 Model', 'L0 conversion model'),
    (r'Discovery/Monitoring Model|monitoring model', 'discovery/monitoring model'),
    (r'Smoothing Model|smoothing', 'smoothing model'),
    (r'Simulator Model|Decision Engine Model|simulator', 'simulator decision engine'),
]

ROLE_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'staff with Chubb LANIDs', 'staff with Chubb LANIDs'),
    (r'business teams?', 'business teams'),
    (r'executives?', 'executives'),
    (r'mobile users?', 'mobile users'),
    (r'desktop users?', 'desktop users'),
    (r'analytics team', 'analytics team'),
    (r'technology users?', 'technology users'),
    (r'investigators?', 'investigators'),
    (r'underwriting operations', 'underwriting operations'),
    (r'IT operations', 'IT operations'),
    (r'pricing analysts?', 'pricing analysts'),
    (r'portfolio managers?', 'portfolio managers'),
]

REGION_SCOPE_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'\bglobal\b', 'Global'),
    (r'\bAPAC\b|Asia[- ]Pacific', 'APAC'),
    (r'\bNA\b|North America', 'North America'),
    (r'\bCOG\b', 'COG'),
    (r'\bEMEA\b|Europe, Middle East', 'EMEA'),
    (r'\bLATAM\b|Latin America', 'LATAM'),
]

DEPLOYMENT_REGION_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'US East 2(?: \(NA Azure\))?', 'US East 2 (NA Azure)'),
    (r'Central US(?: \(NA Azure\))?', 'Central US (NA Azure)'),
    (r'Southeast Asia(?: \(APAC Azure\))?', 'Southeast Asia (APAC Azure)'),
    (r'North Europe(?: \(EMEA Azure\))?', 'North Europe (EMEA Azure)'),
    (r'West Europe(?: \(EMEA Azure\))?', 'West Europe (EMEA Azure)'),
]

CLASSIFICATION_PATTERNS: Sequence[Tuple[str, str]] = [
    (r'Red SPI', 'Red SPI'),
    (r'Red SBI', 'Red SBI'),
    (r'Yellow Business', 'Yellow Business'),
    (r'\bYellow\b', 'Yellow'),
    (r'\bGreen\b', 'Green'),
]

FAMILY_RULES = {
    'Assistants & Document AI': [
        r'document',
        r'question answering',
        r'semantic search',
        r'genai',
        r'openai',
        r'rag',
        r'pathfinder',
        r'chubb ai',
        r'assistant',
        r'comparison',
        r'catalog',
    ],
    'Fraud & Entity Intelligence': [
        r'quantexa',
        r'entity',
        r'fraud',
        r'investigat',
        r'network',
        r'relationship',
        r'graph',
    ],
    'Pricing & Portfolio Optimization': [
        r'pricing',
        r'portfolio',
        r'cvpm',
        r'simulation',
        r'premium',
        r'underwriting',
        r'flow business',
        r'optimization',
    ],
}


def normalize_whitespace(text: str) -> str:
    return re.sub(r'\s+', ' ', text or '').strip()


def canonicalize_heading(text: str) -> str:
    txt = normalize_whitespace(text)
    txt = re.sub(r'^\d+(?:\.\d+)*\s*', '', txt)
    txt = re.sub(r'\s+\d+$', '', txt)
    return txt.strip()


def dedupe_keep_order(items: Iterable[str]) -> List[str]:
    seen = set()
    out: List[str] = []
    for item in items:
        if not item:
            continue
        if item not in seen:
            seen.add(item)
            out.append(item)
    return out


def safe_slug(text: str) -> str:
    return re.sub(r'(^-|-$)', '', re.sub(r'[^a-z0-9]+', '-', text.lower())) or 'app'


def load_doc(path: Path) -> Document:
    return Document(path)


def paragraphs_with_style(doc: Document) -> List[Tuple[str, str]]:
    rows: List[Tuple[str, str]] = []
    for p in doc.paragraphs:
        text = normalize_whitespace(p.text)
        if text:
            rows.append((text, p.style.name if p.style else ''))
    return rows


def document_text(doc: Document) -> str:
    parts: List[str] = []
    for text, _style in paragraphs_with_style(doc):
        parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(normalize_whitespace(cell.text) for cell in row.cells if normalize_whitespace(cell.text))
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


def tables_as_rows(doc: Document) -> List[List[List[str]]]:
    tables: List[List[List[str]]] = []
    for table in doc.tables:
        matrix: List[List[str]] = []
        for row in table.rows:
            matrix.append([normalize_whitespace(cell.text) for cell in row.cells])
        tables.append(matrix)
    return tables


def title_from_doc(doc: Document, filename: str) -> str:
    try:
        raw = doc.tables[0].cell(0, 0).text
        title = normalize_whitespace(raw)
        if title:
            return title
    except Exception:
        pass
    for text, _style in paragraphs_with_style(doc)[:20]:
        if 'Solution Architecture Specification' in text:
            return text
    return filename


def metadata_from_doc(doc: Document) -> Dict[str, str]:
    result: Dict[str, str] = {}
    try:
        table = doc.tables[1]
        for row in table.rows:
            if len(row.cells) >= 2:
                key = normalize_whitespace(row.cells[0].text).rstrip(':')
                val = normalize_whitespace(row.cells[1].text)
                if key:
                    result[key] = val
    except Exception:
        pass
    return result


def split_sections(doc: Document) -> Dict[str, str]:
    rows = paragraphs_with_style(doc)
    idxs: List[Tuple[int, str]] = []
    for idx, (text, style) in enumerate(rows):
        heading = canonicalize_heading(text)
        if heading in SECTION_HEADINGS and (style.startswith('Heading') or heading in SECTION_HEADINGS):
            idxs.append((idx, heading))
    sections: Dict[str, str] = {}
    for pos, (start_idx, heading) in enumerate(idxs):
        end_idx = idxs[pos + 1][0] if pos + 1 < len(idxs) else len(rows)
        body_parts = [rows[i][0] for i in range(start_idx + 1, end_idx)]
        body = ' '.join(part for part in body_parts if part).strip()
        if body:
            sections[heading] = body
    return sections


def extract_si_number(doc: Document, full_text: str) -> str:
    match = re.search(r'SI Number:\s*([A-Za-z0-9\-]+)', full_text)
    if match:
        return match.group(1)
    for table in tables_as_rows(doc):
        if not table:
            continue
        header = ' | '.join(table[0]).lower()
        if 'vers.' in header and 'si' in header:
            for row in table[1:]:
                for idx, cell in enumerate(row):
                    if idx < len(table[0]) and normalize_whitespace(table[0][idx]).lower() == 'si':
                        if cell and cell.upper() != 'XXXX':
                            return cell
    return ''


def section_excerpt(text: str, limit: int = 900) -> str:
    return normalize_whitespace(text)[:limit]


def split_sentences(text: str) -> List[str]:
    base = re.sub(r'\s+', ' ', text)
    base = re.sub(r'(\d+\.)', r' \1 ', base)
    chunks = re.split(r'(?<=[.!?])\s+|\s+\d+\.\s+', base)
    return [normalize_whitespace(chunk) for chunk in chunks if normalize_whitespace(chunk)]


def build_purpose_summary(sections: Dict[str, str]) -> str:
    ordered_sections = ('Background', 'Business Needs and Benefits', 'Objectives and Scope')
    chosen: List[str] = []
    total = 0
    for section_name in ordered_sections:
        section_text = sections.get(section_name, '')
        for sentence in split_sentences(section_text):
            if len(sentence) < 35:
                continue
            if any(sentence.lower() == existing.lower() for existing in chosen):
                continue
            chosen.append(sentence)
            total += len(sentence)
            break
        if len(chosen) >= 3 or total > 550:
            break
    if chosen:
        return ' '.join(chosen)[:700]
    source = ' '.join(sections.get(name, '') for name in ordered_sections if sections.get(name))
    return source[:700]


def find_matches(patterns: Sequence[Tuple[str, str]], corpus: str, limit: int = 12) -> List[str]:
    found: List[str] = []
    for pattern, label in patterns:
        if re.search(pattern, corpus, flags=re.IGNORECASE):
            found.append(label)
    return dedupe_keep_order(found)[:limit]


def find_matches_multi(patterns: Sequence[Tuple[str, str]], corpora: Iterable[str], limit: int = 12) -> List[str]:
    merged = '\n'.join(corpora)
    return find_matches(patterns, merged, limit=limit)


def table_text(table: List[List[str]]) -> str:
    return '\n'.join(' | '.join(cell for cell in row if cell) for row in table)


def parse_security_questionnaire(table: List[List[str]]) -> Dict[str, Dict[str, str]]:
    if not table or len(table[0]) < 4:
        return {}
    header = [c.lower() for c in table[0]]
    if 'question' not in ' '.join(header) or 'answer' not in ' '.join(header):
        return {}
    question_idx = next((i for i, c in enumerate(header) if 'question' in c), None)
    answer_idx = next((i for i, c in enumerate(header) if 'answer' in c), None)
    comments_idx = next((i for i, c in enumerate(header) if 'comments' in c), None)
    if question_idx is None or answer_idx is None:
        return {}
    out: Dict[str, Dict[str, str]] = {}
    for row in table[1:]:
        q = row[question_idx] if question_idx < len(row) else ''
        a = row[answer_idx] if answer_idx < len(row) else ''
        c = row[comments_idx] if comments_idx is not None and comments_idx < len(row) else ''
        qn = normalize_whitespace(q)
        if qn:
            out[qn] = {'answer': normalize_whitespace(a), 'comments': normalize_whitespace(c)}
    return out


def parse_interface_refs(tables: List[List[List[str]]]) -> List[str]:
    refs: List[str] = []
    for table in tables:
        pairs = {normalize_whitespace(row[0]): normalize_whitespace(row[1]) for row in table if len(row) >= 2 and normalize_whitespace(row[0])}
        for key in list(pairs):
            low = key.lower()
            if 'interface reference' in low or low.startswith('attribute name'):
                value = pairs.get(key, '')
                if value:
                    refs.append(re.sub(r'<[^>]+>', '', value))
    return dedupe_keep_order([normalize_whitespace(value) for value in refs if normalize_whitespace(value)])[:12]


def parse_models_from_tables(tables: List[List[List[str]]]) -> List[str]:
    models: List[str] = []
    for table in tables:
        if not table:
            continue
        header = ' | '.join(table[0]).lower()
        if 'model name' in header and 'description' in header:
            for row in table[1:]:
                if row and normalize_whitespace(row[0]):
                    models.append(normalize_whitespace(row[0]))
    return dedupe_keep_order(models)


def parse_rto_rpo(tables: List[List[List[str]]]) -> Dict[str, int | None]:
    result = {'rto_hours': None, 'rpo_hours': None}
    for table in tables:
        for row in table:
            if len(row) >= 2:
                key = normalize_whitespace(row[0]).lower()
                value = normalize_whitespace(row[1])
                if 'recovery time objective' in key:
                    match = re.search(r'(\d+)', value)
                    if match:
                        result['rto_hours'] = int(match.group(1))
                if 'recovery point objective' in key:
                    match = re.search(r'(\d+)', value)
                    if match:
                        result['rpo_hours'] = int(match.group(1))
    return result


def parse_locations_from_tables(tables: List[List[List[str]]], full_text: str) -> List[str]:
    locations: List[str] = []
    for table in tables:
        if not table:
            continue
        header = [cell.lower() for cell in table[0]]
        if 'location' in ' | '.join(header):
            loc_idx = next((i for i, c in enumerate(header) if 'location' in c), None)
            if loc_idx is not None:
                for row in table[1:]:
                    if loc_idx < len(row) and normalize_whitespace(row[loc_idx]):
                        locations.append(normalize_whitespace(row[loc_idx]))
        # Primary/Secondary cloud two-column tables
        for row in table:
            if len(row) >= 2 and normalize_whitespace(row[0]).lower() in {'primary cloud', 'secondary cloud', 'primary data centre', 'secondary data centre'}:
                if normalize_whitespace(row[1]):
                    locations.append(normalize_whitespace(row[1]))
    extracted = []
    corpus = '\n'.join(locations) + '\n' + full_text
    for pattern, label in DEPLOYMENT_REGION_PATTERNS:
        if re.search(pattern, corpus, re.IGNORECASE):
            extracted.append(label)
    return dedupe_keep_order(extracted)


def parse_regions(questionnaire: Dict[str, Dict[str, str]], title: str, filename: str, sections: Dict[str, str]) -> List[str]:
    explicit_chunks = [title, filename]
    purpose_chunks = [sections.get(name, '') for name in ('Background', 'Business Needs and Benefits', 'Objectives and Scope')]
    region_answers = []
    for q, payload in questionnaire.items():
        if 'regions in scope' in q.lower():
            region_answers.append(payload.get('answer', ''))
            region_answers.append(payload.get('comments', ''))
    explicit_joined = '\n'.join(explicit_chunks + region_answers)
    broader_joined = '\n'.join(explicit_chunks + region_answers + purpose_chunks)

    has_explicit_region_answer = any(normalize_whitespace(chunk) for chunk in region_answers)
    regions: List[str] = []

    # Trust explicit answers/title first. GDP/global platform docs default to Global when there is no explicit regional answer.
    if re.search(r'\bglobal\b', explicit_joined, re.IGNORECASE):
        regions.append('Global')
    elif not has_explicit_region_answer and re.search(r'global data platform|\bgdp\b', broader_joined, re.IGNORECASE):
        regions.append('Global')

    for pattern, label in REGION_SCOPE_PATTERNS:
        if label == 'Global':
            continue
        if re.search(pattern, broader_joined, re.IGNORECASE):
            regions.append(label)

    if 'Global' in regions and not has_explicit_region_answer:
        return ['Global']
    return dedupe_keep_order(regions) or ['Global']


def parse_data_classification(questionnaire: Dict[str, Dict[str, str]], interface_tables: List[List[List[str]]]) -> str:
    chunks = []
    for q, payload in questionnaire.items():
        ql = q.lower()
        if 'classification' in ql or 'data stored' in ql or 'data processed' in ql:
            chunks.append(payload.get('answer', ''))
            chunks.append(payload.get('comments', ''))
    joined = '\n'.join(chunks)
    if re.search(r'not currently storing any business data|not currently storing any data|future classifications will be defined', joined, re.IGNORECASE):
        return 'N/A (platform scope)'
    for pattern, label in CLASSIFICATION_PATTERNS:
        if re.search(pattern, joined, re.IGNORECASE):
            return label
    interface_corpus = '\n'.join(table_text(t) for t in interface_tables)
    for pattern, label in CLASSIFICATION_PATTERNS:
        if re.search(pattern, interface_corpus, re.IGNORECASE):
            return label
    return 'Unknown'


def parse_users(questionnaire: Dict[str, Dict[str, str]], full_text: str) -> List[str]:
    chunks = []
    for q, payload in questionnaire.items():
        if 'physical users' in q.lower():
            chunks.append(payload.get('answer', ''))
            chunks.append(payload.get('comments', ''))
    chunks.append(full_text)
    return find_matches_multi(ROLE_PATTERNS, chunks, limit=10)


def parse_review_status(document_status: str, filename: str, title: str) -> str:
    status = (document_status or '').strip().lower()
    filename_lower = filename.lower()
    title_lower = title.lower()
    if 'approved' in status or 'final' in status or 'released' in status:
        return 'Approved'
    if 'gate' in filename_lower or 'gate' in title_lower or 'review' in status:
        return 'In Review'
    return 'Pending Review'


def infer_family(title: str, purpose: str, capabilities: Sequence[str], tech_stack: Sequence[str], models: Sequence[str]) -> str:
    corpus = ' '.join([title, purpose, ' '.join(capabilities), ' '.join(tech_stack), ' '.join(models)]).lower()
    scores = {}
    for family, rules in FAMILY_RULES.items():
        scores[family] = sum(1 for rule in rules if re.search(rule, corpus))
    best_family = max(scores, key=lambda key: scores[key])
    return best_family if scores[best_family] > 0 else 'Other AI Application'


def derive_name(title: str, filename: str) -> Tuple[str, str]:
    clean = title
    clean = re.sub(r'^Solution Architecture Specification\s*', '', clean, flags=re.IGNORECASE).strip()
    clean = re.sub(r'(?<!\s)\(', ' (', clean)
    clean = re.sub(r'\s+', ' ', clean)
    clean = clean or Path(filename).stem
    short = clean
    short = re.sub(r'\(.*?\)', '', short).strip()
    if len(short) > 30:
        short = short[:30].rstrip()
    return clean, short or clean


def extract_business_unit(name: str, title: str, full_text: str) -> str:
    if re.search(r'\bGDP\b', title):
        return 'GDP'
    if re.search(r'\bAPAC\b', title):
        return 'APAC'
    if re.search(r'Chubb AI|Pathfinder', full_text, re.IGNORECASE):
        return 'Enterprise AI'
    return name


def extract_owner(metadata: Dict[str, str], name: str, questionnaire: Dict[str, Dict[str, str]]) -> str:
    author = metadata.get('Author', '')
    return author or name


def priority_from_profile(data_classification: str, review_status: str) -> int:
    if data_classification.startswith('Red'):
        return 5
    if review_status == 'In Review':
        return 4
    if data_classification.startswith('Yellow'):
        return 4
    return 3

def security_answers(questionnaire: Dict[str, Dict[str, str]]) -> str:
    chunks: List[str] = []
    for payload in questionnaire.values():
        chunks.append(payload.get('answer', ''))
        chunks.append(payload.get('comments', ''))
    return '\n'.join(chunk for chunk in chunks if chunk)


def relevant_table_groups(tables: List[List[List[str]]]) -> Dict[str, List[List[List[str]]]]:
    groups = {'interfaces': [], 'models': [], 'technology': [], 'data': [], 'security': []}
    for table in tables:
        text = table_text(table).lower()
        header = ' | '.join(table[0]).lower() if table else ''
        if 'question | answer' in text or ('question' in header and 'answer' in header):
            groups['security'].append(table)
        if 'interface reference' in text or 'attribute name' in text:
            groups['interfaces'].append(table)
            groups['data'].append(table)
        if 'model name' in header and 'description' in header:
            groups['models'].append(table)
        if (
            'components/tiers' in header or
            'cloud function' in header or
            ('product' in header and 'licence' in header) or
            ('location' in header and 'technology type' in header)
        ):
            groups['technology'].append(table)
        if 'source systems' in text or 'input data domains' in text or 'data sources' in text:
            groups['data'].append(table)
    return groups


def source_sections_for(sections: Dict[str, str]) -> Dict[str, List[str]]:
    def present(names: Iterable[str]) -> List[str]:
        return [name for name in names if sections.get(name)]
    return {
        'business_purpose': present(['Background', 'Business Needs and Benefits', 'Objectives and Scope']),
        'capabilities': present(['Logical Architecture', 'System Topology/Context', 'Architecture Layers', 'ML and Analytics', 'ML and Analytics (LLM Models included)']),
        'tech_stack': present(['Architecture Layers', 'Cloud Infrastructure/Services', 'Cloud Software', 'Security Questionnaire']),
        'data_and_security': present(['Introduction', 'Data Dependencies – inputs', 'Data Architecture', 'Security Questionnaire']),
    }


def tokenize(text: str) -> set[str]:
    return {
        token for token in re.sub(r'[^a-z0-9\s]', ' ', text.lower()).split()
        if len(token) > 2 and token not in STOPWORDS
    }


def list_tokens(values: Iterable[str]) -> set[str]:
    merged: set[str] = set()
    for value in values:
        merged |= tokenize(value)
    return merged


def jaccard(a: set[str], b: set[str]) -> float:
    if not a and not b:
        return 0.0
    union = a | b
    if not union:
        return 0.0
    return len(a & b) / len(union)


def pair_score(app_a: Dict, app_b: Dict, weights: Dict[str, float]) -> float:
    purpose = jaccard(tokenize(app_a['business_purpose']), tokenize(app_b['business_purpose']))
    capabilities = jaccard(list_tokens(app_a['capabilities']), list_tokens(app_b['capabilities']))
    features = jaccard(list_tokens(app_a['features']), list_tokens(app_b['features']))
    tech = jaccard(list_tokens(app_a['tech_stack']), list_tokens(app_b['tech_stack']))
    data_domains = jaccard(list_tokens(app_a['data_domains']), list_tokens(app_b['data_domains']))
    region = jaccard(list_tokens(app_a['region_scope']), list_tokens(app_b['region_scope']))
    classification = 1.0 if app_a['data_classification'] == app_b['data_classification'] else 0.0
    data = (0.6 * data_domains) + (0.2 * region) + (0.2 * classification)
    total_weight = sum(weights.values()) or 1.0
    return (
        weights['purpose'] * purpose +
        weights['capabilities'] * capabilities +
        weights['features'] * features +
        weights['tech'] * tech +
        weights['data'] * data
    ) / total_weight


def compute_layout(apps: List[Dict], weights: Dict[str, float], scale: float = 7.0) -> Dict[str, Dict[str, float]]:
    n = len(apps)
    if n == 0:
        return {}
    if n == 1:
        return {apps[0]['id']: {'x': 0.0, 'y': 0.0, 'z': 0.0}}
    if n == 2:
        return {
            apps[0]['id']: {'x': -scale * 0.8, 'y': 0.0, 'z': 0.0},
            apps[1]['id']: {'x': scale * 0.8, 'y': 0.0, 'z': 0.0},
        }
    distances = [[0.0] * n for _ in range(n)]
    for i in range(n):
        for j in range(i + 1, n):
            score = pair_score(apps[i], apps[j], weights)
            dist = max(0.05, 1.0 - score)
            distances[i][j] = distances[j][i] = dist
    mds = MDS(n_components=3, dissimilarity='precomputed', random_state=42, normalized_stress='auto', n_init=4)
    coords = mds.fit_transform(distances)
    max_abs = max(abs(value) for row in coords for value in row) or 1.0
    scaled = coords / max_abs * scale
    return {
        app['id']: {'x': round(float(row[0]), 3), 'y': round(float(row[1]), 3), 'z': round(float(row[2]), 3)}
        for app, row in zip(apps, scaled)
    }


def deep_merge(base: Dict, override: Dict) -> Dict:
    result = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(base.get(key), dict):
            result[key] = deep_merge(base[key], value)
        else:
            result[key] = value
    return result


def load_overrides() -> Dict[str, Dict]:
    if OVERRIDES_PATH.exists():
        return json.loads(OVERRIDES_PATH.read_text(encoding='utf-8'))
    return {}


def build_profile(doc_path: Path, overrides: Dict[str, Dict]) -> Dict:
    doc = load_doc(doc_path)
    full_text = document_text(doc)
    tables = tables_as_rows(doc)
    sections = split_sections(doc)
    metadata = metadata_from_doc(doc)
    title = title_from_doc(doc, doc_path.name)
    name, short_name = derive_name(title, doc_path.name)
    questionnaire = {}
    for table in tables:
        parsed = parse_security_questionnaire(table)
        if parsed:
            questionnaire = parsed
            break

    business_purpose = build_purpose_summary(sections)
    table_groups = relevant_table_groups(tables)
    security_corpus = security_answers(questionnaire)
    architecture_corpus = '\n'.join(filter(None, [
        sections.get('Logical Architecture', ''),
        sections.get('System Topology/Context', ''),
        sections.get('Interface / Integration Item', ''),
        sections.get('Architecture Layers', ''),
        sections.get('ML and Analytics', ''),
        sections.get('ML and Analytics (LLM Models included)', ''),
        sections.get('Objectives and Scope', ''),
        '\n'.join(table_text(t) for t in table_groups['interfaces']),
        '\n'.join(table_text(t) for t in table_groups['models']),
    ]))
    feature_corpus = '\n'.join(filter(None, [
        sections.get('Objectives and Scope', ''),
        sections.get('Interface / Integration Item', ''),
        sections.get('Architecture Layers', ''),
        '\n'.join(table_text(t) for t in table_groups['interfaces']),
    ]))
    tech_corpus = '\n'.join(filter(None, [
        sections.get('Architecture Layers', ''),
        sections.get('Cloud Infrastructure/Services', ''),
        sections.get('Cloud Software', ''),
        sections.get('Introduction', ''),
        security_corpus,
        '\n'.join(table_text(t) for t in table_groups['technology']),
        '\n'.join(table_text(t) for t in table_groups['interfaces']),
        '\n'.join(table_text(t) for t in table_groups['models']),
    ]))
    data_corpus = '\n'.join(filter(None, [
        sections.get('Introduction', ''),
        sections.get('Data Dependencies – inputs', ''),
        sections.get('Data Model', ''),
        security_corpus,
        '\n'.join(table_text(t) for t in table_groups['data']),
    ]))
    capabilities = find_matches(CAPABILITY_PATTERNS, architecture_corpus, limit=10)
    features = find_matches(FEATURE_PATTERNS, feature_corpus, limit=10)
    tech_stack = find_matches(TECH_PATTERNS, tech_corpus, limit=14)
    data_domains = find_matches(DATA_DOMAIN_PATTERNS, data_corpus, limit=12)
    model_names = parse_models_from_tables(tables)
    models = dedupe_keep_order(model_names + find_matches(MODEL_PATTERNS, architecture_corpus + '\n' + security_corpus, limit=10))[:10]
    interfaces = parse_interface_refs(tables)
    region_scope = parse_regions(questionnaire, title, doc_path.name, sections)
    deployment_regions = parse_locations_from_tables(tables, '\n'.join([title, sections.get('Environments and Hosting Location', ''), security_corpus])) or region_scope
    data_classification = parse_data_classification(questionnaire, table_groups['interfaces'])
    users = parse_users(questionnaire, '\n'.join([
        sections.get('User Accessibility', ''),
        sections.get('Background', ''),
        sections.get('Business Needs and Benefits', ''),
        sections.get('Objectives and Scope', ''),
        security_corpus
    ]))
    review_status = parse_review_status(metadata.get('Document Status', ''), doc_path.name, title)
    family = infer_family(title, business_purpose, capabilities, tech_stack, models)
    app_id = safe_slug(name)
    if app_id in {'solution-architecture-specification', ''}:
        app_id = safe_slug(doc_path.stem)
    nfr_nums = parse_rto_rpo(tables)
    availability = sections.get('Availability Hours', '')
    nfr = {
        'rto_hours': nfr_nums['rto_hours'],
        'rpo_hours': nfr_nums['rpo_hours'],
        'availability': availability[:200] if availability else '',
    }

    profile = {
        'id': app_id,
        'name': name,
        'short_name': short_name,
        'family': family,
        'review_status': review_status,
        'owner': extract_owner(metadata, name, questionnaire),
        'business_unit': extract_business_unit(name, title, full_text),
        'region_scope': region_scope,
        'business_purpose': business_purpose,
        'capabilities': capabilities,
        'features': features,
        'tech_stack': tech_stack,
        'data_domains': data_domains,
        'data_classification': data_classification,
        'deployment_regions': deployment_regions,
        'models': models,
        'users': users,
        'interfaces': interfaces,
        'nfr': nfr,
        'priority': priority_from_profile(data_classification, review_status),
        'notes': 'Auto-generated from the real SAS Word document. Add an override only if you want to adjust display labels or enrich the extracted profile.',
        'source_sections': source_sections_for(sections),
        'source_type': 'Real SAS document',
        'document': {
            'filename': doc_path.name,
            'path': f"source_materials/real_sas_docs/{doc_path.name}",
            'title': title,
            'version': metadata.get('Version', ''),
            'author': metadata.get('Author', ''),
            'issue_date': metadata.get('Issue Date', ''),
            'document_status': metadata.get('Document Status', ''),
        },
        'si_number': extract_si_number(doc, full_text),
        'source_excerpt': {
            'background': section_excerpt(sections.get('Background', '')),
            'objectives': section_excerpt(sections.get('Objectives and Scope', '')),
            'architecture_layers': section_excerpt(sections.get('Architecture Layers', '')),
            'security': section_excerpt(sections.get('Security Questionnaire', '')),
        },
    }
    override = overrides.get(doc_path.name) or overrides.get(profile['id']) or {}
    profile = deep_merge(profile, override)
    return profile


def real_doc_paths() -> List[Path]:
    paths = []
    for path in sorted(REAL_SAS_DIR.glob('*.docx')):
        if path.name.startswith('~$'):
            continue
        if 'template' in path.name.lower():
            continue
        paths.append(path)
    return paths


def ensure_unique_profile_ids(profiles: List[Dict]) -> None:
    seen: Dict[str, int] = {}
    for profile in profiles:
        base = profile['id']
        count = seen.get(base, 0)
        seen[base] = count + 1
        if count == 0:
            continue
        suffix = safe_slug(Path(profile['document']['filename']).stem)[:18]
        profile['id'] = f"{base}-{suffix}"


def build_profiles() -> List[Dict]:
    overrides = load_overrides()
    profiles = [build_profile(path, overrides) for path in real_doc_paths()]
    ensure_unique_profile_ids(profiles)

    layouts = {
        'embedding3d': {'purpose': 0.30, 'capabilities': 0.25, 'features': 0.20, 'tech': 0.15, 'data': 0.10},
        'tech_focus': {'purpose': 0.20, 'capabilities': 0.18, 'features': 0.15, 'tech': 0.35, 'data': 0.12},
        'purpose_focus': {'purpose': 0.42, 'capabilities': 0.22, 'features': 0.18, 'tech': 0.10, 'data': 0.08},
    }
    coords_by_layout = {layout_key: compute_layout(profiles, weights) for layout_key, weights in layouts.items()}
    for profile in profiles:
        profile['coords'] = {layout_key: coords[profile['id']] for layout_key, coords in coords_by_layout.items()}
    return profiles


def profile_warnings(profile: Dict) -> List[str]:
    warnings: List[str] = []
    for key in ['business_purpose', 'capabilities', 'features', 'tech_stack', 'data_domains', 'models', 'users']:
        value = profile.get(key)
        if not value:
            warnings.append(f'- Missing or empty `{key}`')
    if not profile.get('si_number'):
        warnings.append('- Missing `si_number`')
    if profile.get('data_classification') == 'Unknown':
        warnings.append('- Could not confidently extract `data_classification`')
    return warnings


def render_extracted_profiles_md(profiles: List[Dict]) -> str:
    lines = [
        '# Real SAS-derived profiles',
        '',
        'These normalized profiles are built dynamically from every `.docx` file in `source_materials/real_sas_docs/`.',
        '',
        'Generation flow:',
        '',
        '1. Scan the real SAS folder for Word documents.',
        '2. Parse section headings, metadata tables, interface tables, hosting tables, model tables, and the security questionnaire.',
        '3. Normalize the extracted values into the workshop profile schema.',
        '4. Compute similarity-driven 3D coordinates and write `applications.json`, `sas_real_profiles.json`, and `assets/workshop-data.js`.',
        '',
    ]
    for app in profiles:
        doc = app['document']
        lines.extend([
            f"## {app['name']}",
            '',
            f"- Source document: `{doc['filename']}`",
            f"- Document title: {doc['title']}",
            f"- Version / issue date: {doc['version']} / {doc['issue_date']}",
            f"- SI number used in workshop: {app['si_number'] or '—'}",
            f"- Family: {app['family']}",
            f"- Business purpose: {app['business_purpose']}",
            f"- Capabilities: {', '.join(app['capabilities']) or '—'}",
            f"- Tech stack: {', '.join(app['tech_stack']) or '—'}",
            f"- Data classification: {app['data_classification']}",
            f"- Interfaces: {', '.join(app['interfaces']) or '—'}",
            '',
        ])
    return '\n'.join(lines)


def render_warnings_md(profiles: List[Dict]) -> str:
    lines = [
        '# Extraction warnings',
        '',
        'These are not hard failures. They simply identify profiles where the generic parser had fewer direct signals and an optional override may help clean up labels for the demo.',
        '',
    ]
    for app in profiles:
        warnings = profile_warnings(app)
        lines.append(f"## {app['name']}")
        if warnings:
            lines.extend(warnings)
        else:
            lines.append('- No notable warnings')
        lines.append('')
    return '\n'.join(lines)


def render_readme(profiles: List[Dict]) -> str:
    profile_lines = '\n'.join(f"- **{app['name']}** — {app['family']}" for app in profiles)
    return f"""# AI Application Landscape Workshop v1.6.1

This package is an **offline interactive demo** for Alan’s AI application landscape and approval workbench.

## What changed in v1.6.1

- The workshop data is now generated **dynamically from every real SAS `.docx` file** in `source_materials/real_sas_docs/`.
- The old hardcoded filename-to-profile map was removed.
- The extractor now parses section headings, interface tables, model tables, hosting tables, and the security questionnaire directly from each Word document.
- An optional `config/profile_overrides.json` file is supported for manual cleanup, but it is **not required** for new files to appear in the demo.
- `scripts/serve_workshop.py` auto-rebuilds the data when files in `real_sas_docs/` change and serves the workshop locally.
- The front-end stays on the v1.3 workshop interaction model, with the additional scatter3d click fix from the attached v1.3 patch (event rebind + native mouseup fallback).
- Fixed a follow-on regression where only the first node click updated the Selected application panel. Plotly handlers are now rebound after `Plotly.react()` finishes, so later clicks keep working.

## Real SAS documents currently included

{profile_lines}

## How the generated files work

- `data/sas_real_profiles.json` — source-of-truth normalized profiles built from the real SAS Word docs.
- `data/applications.json` — the same normalized profiles, kept for the workshop UI.
- `assets/workshop-data.js` — browser-ready copy of the same profile data, used by `index.html` when the workshop runs offline.

These three files are regenerated together by `scripts/extract_real_sas_profiles.py`.

## Dynamic workflow

### Option A — rebuild on demand

1. Drop a new real SAS Word document into `source_materials/real_sas_docs/`
2. Run:

```bash
python scripts/extract_real_sas_profiles.py
```

3. Refresh the workshop.

### Option B — auto-rebuild while serving locally

Run:

```bash
python scripts/serve_workshop.py
```

Then open the local URL that the script prints. When you add or replace a SAS file in `real_sas_docs/`, refresh the browser and the server will rebuild the workshop data automatically.

## Optional overrides

If you want to tweak a label without changing the parser, edit:

- `config/profile_overrides.json`

Example uses:
- rename an app display name
- force a preferred family
- enrich `business_unit` or `owner`
- refine a list such as `capabilities` or `tech_stack`

## Notes

- Direct `index.html` opening still works with the **last generated** data snapshot.
- Auto-detection of newly dropped SAS files requires running `scripts/serve_workshop.py` or rerunning the extractor script.
- This build keeps the v1.3 visual workshop structure rather than the later v1.4 node-trace redesign.
"""


def render_mapping_md() -> str:
    return """# Mapping to SAS template

The normalized workshop profile keeps the same core sections as the SAS template:

| Normalized field | Primary SAS section(s) |
|---|---|
| `business_purpose` | Background, Business Needs and Benefits, Objectives and Scope |
| `capabilities` | Logical Architecture, System Topology/Context, Architecture Layers, ML and Analytics |
| `features` | Objectives and Scope, Interface / Integration Item, Architecture Layers |
| `tech_stack` | Architecture Layers, Cloud Infrastructure/Services, Cloud Software, Security Questionnaire |
| `data_domains` | Data Dependencies – inputs, Data Model, Security Questionnaire |
| `data_classification` | Security Questionnaire |
| `deployment_regions` | Environments and Hosting Location, cloud location tables |
| `models` | ML and Analytics, model tables |
| `interfaces` | Interface / Integration Item, interface tables |
| `users` | User Accessibility, Security Questionnaire |
| `nfr` | Availability Hours, System Failure Conditions |
| `document` | SAS cover/title metadata tables |
"""


def render_demo_script(profiles: List[Dict]) -> str:
    steps = [
        '# Demo script',
        '',
        '1. Run `python scripts/serve_workshop.py` so the workshop auto-rebuilds from the real SAS folder.',
        '2. Open the local workshop URL and let the landscape auto-spin for a few seconds.',
        '3. Click each node to inspect the normalized dimensions extracted from that document.',
        '4. Add a new SAS `.docx` into `source_materials/real_sas_docs/`, refresh the browser, and show the new node appearing after auto-rebuild.',
        '5. Use the lens selector and weight sliders to explain why applications are near each other.',
        '',
        'Current real-document nodes:',
    ]
    steps.extend([f"- {app['name']}" for app in profiles])
    return '\n'.join(steps)


def render_generation_doc() -> str:
    return """# How generation works

The workshop no longer depends on a hardcoded list of known filenames.

## Scan step

The extractor reads every `.docx` file from:

- `source_materials/real_sas_docs/`

It ignores temporary Word lock files such as `~$...docx`.

## Parse step

For each SAS document the extractor reads:

- the cover/title table
- the document metadata table
- heading-based sections such as Background, Objectives and Scope, Architecture Layers, Data Dependencies, ML and Analytics, and Security Questionnaire
- interface tables
- model tables
- cloud location / hosting tables
- the security questionnaire table

## Normalize step

The parser converts those signals into the workshop schema:

- name, id, family, review status
- business purpose
- capabilities
- features
- tech stack
- data domains
- data classification
- deployment regions
- models
- users
- interfaces
- RTO/RPO and availability

## Layout step

After all profiles are built, the script calculates pairwise similarity and derives three 3D coordinate sets:

- `embedding3d`
- `tech_focus`
- `purpose_focus`

## Output step

The extractor writes:

- `data/sas_real_profiles.json`
- `data/applications.json`
- `assets/workshop-data.js`

`workshop-data.js` is simply the browser-ready copy of the same profile list for offline use in `index.html`.

## Optional cleanup

If a new SAS document uses unusual wording, you can refine the generated profile by editing:

- `config/profile_overrides.json`

This is optional. A brand-new SAS document should still appear automatically without adding code.
"""


def file_fingerprint() -> str:
    digest = hashlib.sha256()
    for path in real_doc_paths():
        digest.update(path.name.encode('utf-8'))
        stat = path.stat()
        digest.update(str(stat.st_mtime_ns).encode('utf-8'))
        digest.update(str(stat.st_size).encode('utf-8'))
    if OVERRIDES_PATH.exists():
        stat = OVERRIDES_PATH.stat()
        digest.update(str(stat.st_mtime_ns).encode('utf-8'))
        digest.update(str(stat.st_size).encode('utf-8'))
    return digest.hexdigest()


def write_outputs(profiles: List[Dict]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)

    json_text = json.dumps(profiles, indent=2)
    (DATA_DIR / 'applications.json').write_text(json_text, encoding='utf-8')
    (DATA_DIR / 'sas_real_profiles.json').write_text(json_text, encoding='utf-8')
    (ASSETS_DIR / 'workshop-data.js').write_text('window.WORKSHOP_APPS = ' + json_text + ';\n', encoding='utf-8')
    (DOCS_DIR / 'extracted_profiles.md').write_text(render_extracted_profiles_md(profiles), encoding='utf-8')
    (DOCS_DIR / 'extraction_warnings.md').write_text(render_warnings_md(profiles), encoding='utf-8')
    (DOCS_DIR / 'mapping_to_sas_template.md').write_text(render_mapping_md(), encoding='utf-8')
    (DOCS_DIR / 'demo_script.md').write_text(render_demo_script(profiles), encoding='utf-8')
    (DOCS_DIR / 'how_generation_works.md').write_text(render_generation_doc(), encoding='utf-8')
    (ROOT / 'README.md').write_text(render_readme(profiles), encoding='utf-8')
    (ASSETS_DIR / 'version.txt').write_text(VERSION_TEXT + '\n', encoding='utf-8')
    if not OVERRIDES_PATH.exists():
        OVERRIDES_PATH.write_text('{\n  "_example.docx": {\n    "name": "Friendly display name",\n    "family": "Assistants & Document AI",\n    "business_unit": "Example BU"\n  }\n}\n', encoding='utf-8')


def build_and_write() -> List[Dict]:
    profiles = build_profiles()
    write_outputs(profiles)
    return profiles


def main() -> None:
    parser = argparse.ArgumentParser(description='Build workshop profiles from all real SAS Word documents.')
    parser.parse_args()
    profiles = build_and_write()
    print(f'Built {len(profiles)} real SAS profile(s) from {REAL_SAS_DIR}')


if __name__ == '__main__':
    main()
